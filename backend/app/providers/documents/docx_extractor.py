from __future__ import annotations

import io

from app.models.schemas import ExtractedDocument


def extract_docx(file_bytes: bytes) -> ExtractedDocument:
    try:
        import docx
    except ImportError:
        return ExtractedDocument(structure_notes=["python-docx not installed; cannot parse document."])

    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables_text = []
    for table in document.tables:
        for row in table.rows:
            tables_text.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(paragraphs + tables_text)
    notes = [f"{len(document.tables)} table(s) detected."] if document.tables else []
    return ExtractedDocument(text=text, pages=1, used_ocr=False, structure_notes=notes)
